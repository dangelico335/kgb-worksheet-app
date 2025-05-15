
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import tempfile

app = Flask(__name__)
supported_files = ['A#m_or_Bbm_Guitar.png', 'A#m_or_Bbm_Piano.png', 'A#_or_Bb_Bass.png', 'A#_or_Bb_Guitar.png', 'A#_or_Bb_Piano.png', 'Am_Guitar.png', 'Am_Piano.png', 'A_Bass.png', 'A_Guitar.png', 'A_Piano.png', 'Bm_Guitar.png', 'Bm_Piano.png', 'B_Bass.png', 'B_Guitar.png', 'B_Piano.png', 'C#m_or_Dbm_Guitar.png', 'C#m_or_Dbm_Piano.png', 'C#_or_Db_Bass.png', 'C#_or_Db_Guitar.png', 'C#_or_Db_Piano.png', 'Cm_Guitar.png', 'Cm_Piano.png', 'C_Bass.png', 'C_Guitar.png', 'C_Piano.png', 'D#m_or_Ebm_Guitar.png', 'D#m_or_Ebm_Piano.png', 'D#_or_Eb_Bass.png', 'D#_or_Eb_Guitar.png', 'D#_or_Eb_Piano.png', 'Dm_Guitar.png', 'Dm_Piano.png', 'D_Bass.png', 'D_Guitar.png', 'D_Piano.png', 'Em_Guitar.png', 'Em_Piano.png', 'E_Bass.png', 'E_Guitar.png', 'E_Piano.png', 'F#m_or_Gbm_Guitar.png', 'F#m_or_Gbm_Piano.png', 'F#_or_Gb_Bass.png', 'F#_or_Gb_Guitar.png', 'F#_or_Gb_Piano.png', 'Fm_Guitar.png', 'Fm_Piano.png', 'F_Bass.png', 'F_Guitar.png', 'F_Piano.png', 'G#m_or_Abm_Guitar.png', 'G#m_or_Abm_Piano.png', 'G#_or_Ab_Bass.png', 'G#_or_Ab_Guitar.png', 'G#_or_Ab_Piano.png', 'Gm_Guitar.png', 'Gm_Piano.png', 'G_Bass.png', 'G_Guitar.png', 'G_Piano.png']

def get_image_filename(chord, instrument):
    root = chord.replace("add9", "").strip()
    if instrument == "Bass":
        root = root.replace("m", "")
    for f in supported_files:
        if f.startswith(root) and instrument in f:
            return f"static/images/{f}"
    return None

def add_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for edge in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcPr.append(border)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        title = request.form['title']
        composer = request.form['composer']
        key = request.form['key']
        instruments = request.form.getlist('instruments')

        sections = []
        for i in range(1, 4):
            name = request.form.get(f'section{i}_name')
            chords = request.form.get(f'section{i}_chords')
            if name and chords:
                chord_list = [c.strip() for c in chords.split(',')]
                sections.append((name, chord_list))

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = "Bodoni MT Black"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bodoni MT Black")
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        composer_para = doc.add_paragraph()
        composer_run = composer_para.add_run(composer)
        composer_run.font.name = "Bodoni MT Black"
        composer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bodoni MT Black")
        composer_run.font.size = Pt(9)
        composer_run.font.color.rgb = RGBColor(0, 0, 0)
        composer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        for i, (section_name, chord_list) in enumerate(sections):
            if i > 0:
                doc.add_page_break()
            doc.add_paragraph().add_run(section_name).bold = True
            chord_groups = [[c] if "/" not in c else c.split("/") for c in chord_list]
            flat_chords = [ch for group in chord_groups for ch in group]
            table = doc.add_table(rows=len(instruments)+1, cols=len(flat_chords)+1)
            table.cell(0, 0).text = ""
            col_index = 1
            for group in chord_groups:
                label = " / ".join(group)
                para = table.cell(0, col_index).paragraphs[0]
                para.add_run(label).bold = True
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if len(group) == 2:
                    table.cell(0, col_index).merge(table.cell(0, col_index+1))
                    col_index += 2
                else:
                    col_index += 1
            for idx, instr in enumerate(instruments):
                row = table.rows[idx+1]
                row.cells[0].text = instr
                col_index = 1
                for group in chord_groups:
                    for chord in group:
                        para = row.cells[col_index].paragraphs[0]
                        path = get_image_filename(chord, instr)
                        if path and os.path.exists(path):
                            para.add_run().add_picture(path, width=Inches(1.2))
                        else:
                            para.add_run("[Missing]")
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        col_index += 1
            for row in table.rows:
                for cell in row.cells:
                    add_borders(cell)

        doc.add_paragraph()
        doc.add_paragraph(f"Key: {key}").bold = True
        for _ in range(4):
            doc.add_paragraph("_________________________")

        tmp_path = tempfile.mktemp(suffix=".docx")
        doc.save(tmp_path)
        return send_file(tmp_path, as_attachment=True)

    return render_template("index.html")
